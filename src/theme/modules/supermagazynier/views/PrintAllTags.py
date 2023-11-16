import base64
from django.db import connections
from rest_framework.views import APIView
from rest_framework.response import Response
import os
import shutil
import zipfile
from PIL import Image
import zpl
from docx import Document
from docx.shared import Mm
from django.conf import settings
from django.http import HttpResponse, Http404
from theme.modules.supermagazynier.views.Zpl2Generator import Zpl2Generator
from theme.modules.supermagazynier.models import Document as DocumentObject, ListOfProducts as ListOfProductsObject
from theme.modules.supermagazynier.serializers import DocumentSerializer, ListOfProductsSerializer

class PrintAllTags(APIView):
    def get(self, request, document_id, filename):
        lop = self.get_product_atributes(document_id)
        self.generate_docx_with_labels(lop, filename)
        response = self.download("/home/grzegorz/metalzbyt_panel/src/labels.docx")
        return HttpResponse(response, content_type="application/msword")
    
    def download(self, path):
        file_path = os.path.join(settings.MEDIA_ROOT, path)
        if os.path.exists(file_path):
            with open(file_path, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type="application/msword")
                response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_path)
                return response
        raise Http404
    
    def get_product_atributes(self, document_id):
        list_of_products = self.getListOfProductsFromDocument(document_id)
        return list_of_products
    
    def generate_label_image(self, not_formatted_product, filename):
        product = {
            'code': str(not_formatted_product['Product']['Kod']),
            'EAN': str(not_formatted_product['Product']['EAN']),
            'name': str(not_formatted_product['Product']['Nazwa']),
            'price': str(format(not_formatted_product['Product']['CenaDetaliczna']*1.23, '.2f')),
            'jm': str(not_formatted_product['Product']['jm']),
            'quantity': not_formatted_product['Quantity']
        }
        Generator = Zpl2Generator(product)
        readed_file = Generator.read_file(filename)
        image = readed_file[1]
        image_path = "/home/grzegorz/metalzbyt_panel/src/theme/modules/supermagazynier/assets/preview.png"
        image.save(image_path,"png")
        encoded_string = ''
        with open(image_path, 'rb') as image_file:
            encoded_string = base64.b64encode(image_file.read())
        width, height = image.size
        return {'image': encoded_string, 'size': [width, height]}
    
    def generate_docx_with_labels(self, lop, filename):
        document = Document()
        section = document.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(5)
        section.right_margin = Mm(5)
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        section.header_distance = Mm(0)
        section.footer_distance = Mm(0)
        section.left_indent = Mm(0)
        
        p = document.add_paragraph()
        p.style = 'No Spacing'
        r = p.add_run()
        image_path = "/home/grzegorz/metalzbyt_panel/src/theme/modules/supermagazynier/assets/preview.png"
        for product in lop:
            image = self.generate_label_image(product, filename)
            for i in range(int(product['Quantity'])):
                r.add_picture(image_path, width=Mm(image['size'][0]/7.98))
        
        document.save('labels.docx')
        self.delete_spacing_from_elements()
        
    def delete_spacing_from_elements(self):
        WORKING_DIR = os.getcwd()
        TEMP_DOCX = os.path.join(WORKING_DIR, "/home/grzegorz/metalzbyt_panel/src/labels.docx")
        TEMP_ZIP = os.path.join(WORKING_DIR, "/home/grzegorz/metalzbyt_panel/src/labels.zip")
        TEMP_FOLDER = os.path.join(WORKING_DIR, "/home/grzegorz/metalzbyt_panel/src/labels")

        # remove old zip file or folder template
        if os.path.exists(TEMP_ZIP):
            os.remove(TEMP_ZIP)
        if os.path.exists(TEMP_FOLDER):
            shutil.rmtree(TEMP_FOLDER)

        # reformat template.docx's extension
        os.rename(TEMP_DOCX, TEMP_ZIP)

        # unzip file zip to specific folder
        with zipfile.ZipFile(TEMP_ZIP, 'r') as z:
            z.extractall(TEMP_FOLDER)

        # change header xml file
        header_xml = os.path.join(TEMP_FOLDER, "word", "document.xml")
        xmlstring = open(header_xml, 'r', encoding='utf-8').read()
        xmlstring = xmlstring.replace("inline ", "inline distT='0' distB='0' distL='0' distR='0' ")
        with open(header_xml, "wb") as f:
            f.write(xmlstring.encode("UTF-8"))

        # zip temp folder to zip file
        os.remove(TEMP_ZIP)
        shutil.make_archive(TEMP_ZIP.replace(".zip", ""), 'zip', TEMP_FOLDER)

        # rename zip file to docx
        os.rename(TEMP_ZIP, TEMP_DOCX)
        shutil.rmtree(TEMP_FOLDER)

    def getListOfProductsFromDocument(self, pk):
        list_of_products = ListOfProductsObject.objects.filter(Document_id=pk)
        lop = []
        for p in list_of_products:
            with connections['METALZBYT_KOLEKTOR'].cursor() as cursor:
                sql = f'''with tabela as  
                        (select  ROW_NUMBER() over(order by T0.kod asc) as RowNumber,
                        T9.Data as 'maincategory',T10.Data as 'producent', 
                        T3.Data as active,T0.ID, T0.Guid,T0.EAN, T0.NumerKatalogowy,T2.Kod as jm, T0.Kod, T0.Nazwa,T0.Opis,  
                        (SELECT sum(T1.[IloscValue]) FROM Zasoby T1  
                        WHERE T1.Magazyn != 8 AND T0.id = T1.Towar AND T1.PartiaPierwotnaTyp = 1  GROUP BY T1.Towar) as 'quantity', 
                        (SELECT MAX(T8.PartiaWartosc/T8.IloscValue) FROM Zasoby T8  
                        WHERE T0.id = T8.Towar AND T8.PartiaPierwotnaTyp = 1 GROUP BY T8.Towar) as 'cenazakupu', 
                        T5.NettoValue as 'Cena hurtowa', T6.NettoValue as 'Cena detaliczna', 
                        T7.NettoValue as 'Cena sklep internetowy', (SELECT sum(T1.[IloscValue]) FROM Zasoby T1  WHERE T1.Magazyn = 9 AND T0.id = T1.Towar AND T1.PartiaPierwotnaTyp = 1  GROUP BY T1.Towar) as m2,(SELECT sum(T1.[IloscValue]) FROM Zasoby T1  WHERE T1.Magazyn = 5 AND T0.id = T1.Towar AND T1.PartiaPierwotnaTyp = 1  GROUP BY T1.Towar) as m1,(SELECT sum(T1.[IloscValue]) FROM Zasoby T1  WHERE T1.Magazyn = 1 AND T0.id = T1.Towar AND T1.PartiaPierwotnaTyp = 1  GROUP BY T1.Towar) as m6
                        FROM Towary T0  
                        LEFT JOIN Jednostki T2 ON T2.ID = T0.Jednostka  
                        LEFT JOIN Features T3 ON (T3.ParentType LIKE 'Towary' 
                        AND T3.[Name] LIKE 'Sklep internetowy' 
                        AND T3.Parent=T0.ID) 
                        LEFT JOIN Features T10 ON (T10.Parent=T0.ID AND T10.ParentType LIKE 'Towary' AND T10.[Name] LIKE 'Producent') 
                        LEFT JOIN  Features T9 ON (T9.Parent=T0.ID AND T9.ParentType LIKE 'Towary' AND T9.[Name] LIKE 'Asortyment')  
                        LEFT JOIN  Ceny T5 ON (T5.Towar = T0.ID AND T5.Definicja = 2 AND T5.NettoSymbol='PLN')  
                        LEFT JOIN  Ceny T6 ON (T6.Towar = T0.ID AND T6.Definicja = 3 AND T6.NettoSymbol='PLN')  
                        LEFT JOIN  Ceny T7 ON (T7.Towar = T0.ID AND T7.Definicja = 9 AND T7.NettoSymbol='PLN') 
                        WHERE T0.id LIKE '{p.Product}') 
                        select * from tabela where RowNumber  BETWEEN 0 AND 10000000'''
                cursor.execute(sql)
                row = cursor.fetchone()
                jsonRow = {
                    'lp': row[0],
                    'Kategoria': row[1],
                    'Producent': row[2],
                    'cos': row[3],
                    'id': row[4],
                    'Guid': row[5],
                    'EAN': row[6],
                    'NumerKatalogowy': row[7],
                    'jm': row[8],
                    'Kod': row[9],
                    'Nazwa': row[10],
                    'Opis': row[11],
                    'Ilosc': row[12],
                    'CenaZakupu': row[13],
                    'CenaDetaliczna': row[14],
                    'CenaSklepInternetowy': row[15],
                    'cos2': row[16],
                    'm2': row[17],
                    'm1': row[18],
                    'm6': row[19]
                }
                obj = {
                    'Document': p.Document_id,
                    'Product': jsonRow,
                    'Quantity': p.Quantity
                }
                lop.append(obj)
        return lop